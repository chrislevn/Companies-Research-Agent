"""Three renderings of one brief: Markdown is canonical, HTML for the webapp,
PDF for the thing a person actually forwards.

Both obey the same rule — an unverified claim is *marked*, never hidden, and
never silently dropped. A reader who cannot tell which lines are evidenced will
assume all of them are, which is exactly the failure a sourced brief exists to
prevent.
"""

from __future__ import annotations

import logging
import platform
from html import escape
from pathlib import Path

from ..models import Brief, BriefClaim
from .build import PROFILE_FIELDS

log = logging.getLogger(__name__)

LABELS = dict(PROFILE_FIELDS) | {
    "contact_name": "Contact",
    "contact_title": "Title",
    "intent": "What they want",
    "relationship": "Relationship",
}


def _label(field: str) -> str:
    if field.startswith("news["):
        return "News"
    if field.startswith("meeting_prep["):
        return "Prep"
    return LABELS.get(field, field.replace("_", " ").capitalize())


def _grouped(brief: Brief) -> tuple[list[BriefClaim], list[BriefClaim], list[BriefClaim]]:
    """Facts, news and prep — the three things a reader wants separately."""
    facts, news, prep = [], [], []
    for claim in brief.claims:
        if claim.field.startswith("news["):
            news.append(claim)
        elif claim.field.startswith("meeting_prep["):
            prep.append(claim)
        else:
            facts.append(claim)
    return facts, news, prep


# --- markdown (canonical) --------------------------------------------------


def to_markdown(brief: Brief) -> str:
    facts, news, prep = _grouped(brief)
    out: list[str] = [f"# {brief.company}", ""]
    if brief.domain:
        out.append(f"**{brief.domain}** · generated {brief.generated_at:%Y-%m-%d %H:%M} UTC")
        out.append("")

    if brief.upcoming_meeting:
        meeting = brief.upcoming_meeting
        out += [
            "## Upcoming meeting",
            "",
            f"**{meeting.starts_at:%A %d %B %Y, %H:%M}** — {meeting.title or '(untitled)'}",
            "",
            f"Matched on {meeting.matched_on.replace('_', ' ')} "
            f"(confidence {meeting.confidence:.2f}).",
            "",
        ]
        if meeting.attendees:
            out += ["Attendees: " + ", ".join(meeting.attendees), ""]

    if facts:
        out += ["## The company", ""]
        for claim in facts:
            out.append(f"- **{_label(claim.field)}:** {claim.value}{_md_source(claim)}")
        out.append("")

    if brief.talking_points:
        out += ["## Talking points", ""]
        out += [f"- {point}" for point in brief.talking_points]
        out += ["", "*Only points traceable to a source appear here.*", ""]

    unsourced_prep = [c for c in prep if not c.verified]
    if unsourced_prep:
        out += ["## Suggested prep (unverified)", ""]
        for claim in unsourced_prep:
            out.append(f"- {claim.value} *(no source — treat as a prompt, not a fact)*")
        out.append("")

    if news:
        out += ["## Recent news", ""]
        for claim in news:
            out.append(f"- {claim.value}{_md_source(claim)}")
        out.append("")

    if brief.unknowns:
        out += ["## What we could not establish", ""]
        out += [f"- {gap}" for gap in brief.unknowns]
        out += ["", "*Listed rather than left blank: a gap you cannot see is a gap "
                "you will fill in yourself.*", ""]

    if brief.sources:
        out += ["## Sources", ""]
        out += [f"{i}. <{url}>" for i, url in enumerate(brief.sources, start=1)]
        out.append("")

    out += ["---", "", f"Status: **{brief.status}** · lead `{brief.lead_id}`"]
    if brief.approved_by:
        out.append(f" · approved by {brief.approved_by}")
    return "\n".join(out).rstrip() + "\n"


def _md_source(claim: BriefClaim) -> str:
    if not claim.source_url:
        return "  *(unverified)*"
    if claim.source_url.startswith("email:"):
        return "  *(from the email)*"
    return f"  ([source]({claim.source_url}))"


# --- html (webapp) ---------------------------------------------------------


def to_html(brief: Brief) -> str:
    facts, news, prep = _grouped(brief)
    parts: list[str] = [
        f'<article class="brief" data-status="{escape(brief.status)}">',
        f"<header><h1>{escape(brief.company)}</h1>",
    ]
    if brief.domain:
        parts.append(
            f'<p class="muted small">{escape(brief.domain)} · generated '
            f"{brief.generated_at:%Y-%m-%d %H:%M} UTC</p>"
        )
    parts.append("</header>")

    if brief.upcoming_meeting:
        meeting = brief.upcoming_meeting
        parts += [
            '<section class="meeting"><h2>Upcoming meeting</h2>',
            f"<p><strong>{meeting.starts_at:%A %d %B %Y, %H:%M}</strong> — "
            f"{escape(meeting.title or '(untitled)')}</p>",
            f'<p class="muted small">Matched on '
            f"{escape(meeting.matched_on.replace('_', ' '))} "
            f"(confidence {meeting.confidence:.2f})</p>",
        ]
        if meeting.attendees:
            parts.append(
                '<p class="muted small">Attendees: '
                + escape(", ".join(meeting.attendees)) + "</p>"
            )
        parts.append("</section>")

    if facts:
        parts.append("<section><h2>The company</h2><dl>")
        for claim in facts:
            parts.append(
                f"<dt>{escape(_label(claim.field))}</dt>"
                f"<dd>{escape(claim.value or '')}{_html_source(claim)}</dd>"
            )
        parts.append("</dl></section>")

    if brief.talking_points:
        parts.append("<section><h2>Talking points</h2><ul>")
        parts += [f"<li>{escape(point)}</li>" for point in brief.talking_points]
        parts.append('</ul><p class="muted small">Only points traceable to a source '
                     "appear here.</p></section>")

    unsourced_prep = [c for c in prep if not c.verified]
    if unsourced_prep:
        parts.append('<section class="unverified"><h2>Suggested prep (unverified)</h2><ul>')
        for claim in unsourced_prep:
            parts.append(
                f"<li>{escape(claim.value or '')} "
                '<span class="flag">no source — treat as a prompt, not a fact</span></li>'
            )
        parts.append("</ul></section>")

    if news:
        parts.append("<section><h2>Recent news</h2><ul>")
        for claim in news:
            parts.append(f"<li>{escape(claim.value or '')}{_html_source(claim)}</li>")
        parts.append("</ul></section>")

    if brief.unknowns:
        parts.append('<section class="unknowns"><h2>What we could not establish</h2><ul>')
        parts += [f"<li>{escape(gap)}</li>" for gap in brief.unknowns]
        parts.append('</ul><p class="muted small">Listed rather than left blank: a gap '
                     "you cannot see is a gap you will fill in yourself.</p></section>")

    if brief.sources:
        parts.append("<section><h2>Sources</h2><ol>")
        for url in brief.sources:
            safe = escape(url, quote=True)
            parts.append(
                f'<li><a href="{safe}" target="_blank" rel="noopener">{escape(url)}</a></li>'
            )
        parts.append("</ol></section>")

    parts.append(
        f'<footer class="muted small">Status: <strong>{escape(brief.status)}</strong> · '
        f"lead <code>{escape(brief.lead_id)}</code></footer></article>"
    )
    return "\n".join(parts)


def _html_source(claim: BriefClaim) -> str:
    if not claim.source_url:
        return ' <span class="flag">unverified</span>'
    if claim.source_url.startswith("email:"):
        return ' <span class="muted small">from the email</span>'
    safe = escape(claim.source_url, quote=True)
    return f' <a class="src" href="{safe}" target="_blank" rel="noopener">source</a>'


# --- pdf --------------------------------------------------------------------

# A PDF's built-in fonts are Latin-1. This agent reads Vietnamese mail, so a
# built-in font would turn "Công ty Cổ phần Điện Thủ Đức" into punctuation — and
# do it silently, in the one artefact that gets forwarded to other people. So a
# Unicode TrueType face is resolved from the system, in the order a platform is
# likeliest to have one, and the renderer says so in the log when it cannot.
FONT_CANDIDATES = {
    "Darwin": [
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ],
    "Windows": [
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/segoeui.ttf",
        "C:/Windows/Fonts/tahoma.ttf",
    ],
    "Linux": [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
        "/usr/share/fonts/TTF/DejaVuSans.ttf",
    ],
}


class PdfUnavailable(RuntimeError):
    """fpdf2 is not installed. The other renderings still work."""


def _unicode_font() -> tuple[Path, Path | None] | None:
    """Regular and bold faces, or None if the machine has no usable font."""
    for candidate in FONT_CANDIDATES.get(platform.system(), []):
        path = Path(candidate)
        if not path.exists():
            continue
        bold = path.with_name(path.stem + " Bold" + path.suffix)
        if not bold.exists():
            bold = path.with_name(path.stem.replace("-Regular", "-Bold") + path.suffix)
        return path, (bold if bold.exists() else None)
    return None


# The app's own palette, so a printed brief and the screen it came from are
# recognisably the same document.
_INK = (26, 26, 24)
_MUTED = (107, 107, 102)
_LINE = (226, 226, 221)
_ACCENT = (180, 85, 31)
_ACCENT_SOFT = (253, 241, 232)
_WARN = (138, 90, 5)
_OK = (31, 122, 77)

_STATUS_COLOURS = {
    "approved": _OK,
    "delivered": _OK,
    "rejected": (179, 38, 30),
    "draft": _MUTED,
}


def _source_note(claim: BriefClaim) -> tuple[str, bool]:
    """How a claim's provenance reads, and whether it is a warning.

    Shared by the PDF and the Word file so all three renderings use one
    vocabulary. They diverged once: the exports said "not verified against a
    source" where the Markdown said "unverified", and printed the raw
    `email:<id>` sentinel as though it were a URL a reader could open.
    """
    url = claim.source_url or ""
    if not url:
        return "unverified", True
    if url.startswith("email:"):
        return "from the email", False
    return url, False


def _pdf_class(family: str, brief: Brief):
    """A document with a running header and a numbered footer.

    Subclassed rather than drawn inline because fpdf2 calls ``header`` and
    ``footer`` on every page break, which is the only way a multi-page brief
    keeps its identity on page 3. A reader who has printed six of these needs to
    know which company page 3 belongs to without hunting for page 1.
    """
    from fpdf import FPDF

    class _BriefDoc(FPDF):
        def header(self) -> None:
            # Page 1 carries the masthead instead; a running header above it
            # would just repeat what the title already says.
            if self.page_no() == 1:
                return
            self.set_y(10)
            self.set_font(family, "", 8)
            self.set_text_color(*_MUTED)
            self.cell(0, 4, (brief.company or "Company brief")[:70])
            self.set_font(family, "", 8)
            self.cell(0, 4, "COMPANY BRIEF", align="R",
                      new_x="LMARGIN", new_y="NEXT")
            self.ln(1)
            self.set_draw_color(*_LINE)
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
            self.ln(4)

        def footer(self) -> None:
            self.set_y(-14)
            self.set_draw_color(*_LINE)
            self.line(self.l_margin, self.get_y(), self.w - self.r_margin, self.get_y())
            self.ln(2)
            self.set_font(family, "", 7.5)
            self.set_text_color(*_MUTED)
            left = f"{brief.domain or brief.company or ''} · " \
                   f"generated {brief.generated_at:%Y-%m-%d %H:%M} UTC"
            self.cell(0, 4, left[:88])
            self.cell(0, 4, f"{self.page_no()} / {{nb}}", align="R")

    return _BriefDoc




def to_pdf(brief: Brief) -> bytes:
    """The brief as a designed PDF — the artefact a person forwards.

    Unverified claims stay marked here exactly as they are in the Markdown. A
    PDF looks more official, which makes dropping the caveat more tempting and
    more dangerous, not less: the whole value of the document is that a reader
    can see which lines are evidenced.
    """
    try:
        from fpdf import FPDF  # noqa: F401  (import guard only)
    except Exception as exc:  # pragma: no cover - import guard
        raise PdfUnavailable(f"fpdf2 is not installed: {exc}") from exc

    facts, news, prep = _grouped(brief)

    faces = _unicode_font()
    family = "body" if faces else "Helvetica"
    pdf = _pdf_class(family, brief)(unit="mm", format="A4")
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_margins(18, 14, 18)
    if faces:
        regular, bold = faces
        pdf.add_font("body", "", str(regular))
        pdf.add_font("body", "B", str(bold) if bold else str(regular))
    else:
        log.warning(
            "No Unicode font found on this system; the PDF falls back to Helvetica "
            "and non-Latin text (Vietnamese, in particular) will not render."
        )
    pdf.add_page()
    width = pdf.w - pdf.l_margin - pdf.r_margin

    def text(body: str, *, size: float = 10, style: str = "", colour=_INK,
             gap: float = 5.0, indent: float = 0.0) -> None:
        pdf.set_font(family, style, size)
        pdf.set_text_color(*colour)
        if indent:
            pdf.set_x(pdf.l_margin + indent)
        pdf.multi_cell(width - indent, gap, body, align="L",
                       new_x="LMARGIN", new_y="NEXT")

    def rule(colour=_LINE, thickness: float = 0.2, space_before: float = 3.0,
             space_after: float = 3.0) -> None:
        pdf.ln(space_before)
        pdf.set_draw_color(*colour)
        pdf.set_line_width(thickness)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.set_line_width(0.2)
        pdf.ln(space_after)

    def heading(label: str) -> None:
        # Keep a heading with at least one line of what follows.
        if pdf.get_y() > pdf.h - 46:
            pdf.add_page()
        pdf.ln(3.5)
        pdf.set_font(family, "B", 8)
        pdf.set_text_color(*_ACCENT)
        pdf.cell(0, 4, " ".join(label.upper()), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(0.5)
        pdf.set_draw_color(*_LINE)
        pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
        pdf.ln(3)

    # --- masthead ----------------------------------------------------------
    pdf.set_fill_color(*_ACCENT)
    pdf.rect(pdf.l_margin, pdf.get_y(), 26, 1.6, style="F")
    pdf.ln(6)
    text("COMPANY BRIEF", size=8, style="B", colour=_ACCENT, gap=4)
    pdf.ln(1)
    text(brief.company or "(unknown company)", size=22, style="B", gap=10)

    meta = " · ".join(filter(None, [
        brief.domain,
        f"generated {brief.generated_at:%d %B %Y, %H:%M} UTC",
    ]))
    text(meta, size=9, colour=_MUTED, gap=5)

    # Status reads as a pill rather than another line of grey text: it is the
    # first thing a reader checks and the only field that changes after issue.
    status = (brief.status or "draft").lower()
    colour = _STATUS_COLOURS.get(status, _MUTED)
    pdf.set_font(family, "B", 7.5)
    label = status.upper()
    pill = pdf.get_string_width(label) + 6
    pdf.set_fill_color(*_ACCENT_SOFT if status in ("approved", "delivered") else (245, 245, 243))
    pdf.set_text_color(*colour)
    pdf.cell(pill, 5.5, label, align="C", fill=True, new_x="LMARGIN", new_y="NEXT")
    if brief.approved_by:
        stamp = f"approved by {brief.approved_by}"
        if brief.approved_at:
            stamp += f" on {brief.approved_at:%d %B %Y, %H:%M} UTC"
        pdf.ln(1.5)
        text(stamp, size=8.5, colour=_MUTED, gap=4)
    rule(space_before=4, space_after=1)

    # --- meeting -----------------------------------------------------------
    if brief.upcoming_meeting:
        meeting = brief.upcoming_meeting
        heading("Upcoming meeting")
        top = pdf.get_y()
        pdf.set_fill_color(*_ACCENT_SOFT)
        pdf.rect(pdf.l_margin, top - 1.5, width, 13, style="F")
        text(f"{meeting.starts_at:%A %d %B %Y, %H:%M}  —  "
             f"{meeting.title or '(untitled)'}", size=10.5, style="B", gap=5.5,
             indent=3)
        if getattr(meeting, "attendees", None):
            text(", ".join(meeting.attendees)[:110], size=8.5, colour=_MUTED,
                 gap=4, indent=3)
        pdf.ln(2)

    # --- claims ------------------------------------------------------------
    def claim_block(title: str, claims: list[BriefClaim]) -> None:
        if not claims:
            return
        heading(title)
        for claim in claims:
            pdf.set_font(family, "B", 9.5)
            pdf.set_text_color(*_INK)
            label = f"{_label(claim.field)}:  "
            pdf.cell(pdf.get_string_width(label), 5, label)
            pdf.set_font(family, "", 9.5)
            pdf.multi_cell(width - pdf.get_string_width(label), 5,
                           str(claim.value), align="L",
                           new_x="LMARGIN", new_y="NEXT")
            note, warn = _source_note(claim)
            text(note, size=7.5, colour=_WARN if warn else _MUTED,
                 gap=3.6, indent=3)
            pdf.ln(1.2)

    claim_block("What we know", facts)
    claim_block("Recent news", news)
    claim_block("Worth raising", prep)

    if brief.talking_points:
        heading("Talking points")
        for point in brief.talking_points:
            text(f"—  {point}", size=9.5, gap=5)
            pdf.ln(0.8)

    if brief.unknowns:
        heading("Not established")
        for unknown in brief.unknowns:
            text(f"—  {unknown}", size=9.5, colour=_MUTED, gap=5)
            pdf.ln(0.8)

    if brief.sources:
        heading("Sources")
        for n, url in enumerate(brief.sources, 1):
            text(f"{n}.  {url}", size=7.5, colour=_MUTED, gap=3.8)

    out = pdf.output()
    return bytes(out)


# --- word -------------------------------------------------------------------


class DocxUnavailable(RuntimeError):
    """python-docx is not installed. The other renderings still work."""


def to_docx(brief: Brief) -> bytes:
    """The brief as an editable .docx.

    The PDF is the artefact you forward; this is the one you *change* before
    forwarding — a salesperson adding what they know that the agent could not
    find. So this is built from real Word styles (Title, Heading 1, List
    Bullet) rather than hand-set fonts: styled text keeps working when someone
    edits it, restyles the document, or pastes it into their own template,
    whereas direct formatting survives none of those.

    Source lines stay attached to their claims, and unverified claims stay
    marked. An editable brief is exactly where a caveat is most likely to be
    quietly deleted, so it is written in as text rather than as a colour a
    reader might not notice they removed.
    """
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt, RGBColor
    except Exception as exc:  # pragma: no cover - import guard
        raise DocxUnavailable(f"python-docx is not installed: {exc}") from exc

    facts, news, prep = _grouped(brief)
    doc = Document()

    def muted(text_value: str, *, size: float = 8.5, italic: bool = False,
              colour=(107, 107, 102)) -> None:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        run = paragraph.add_run(text_value)
        run.font.size = Pt(size)
        run.italic = italic
        run.font.color.rgb = RGBColor(*colour)

    eyebrow = doc.add_paragraph()
    eyebrow.paragraph_format.space_after = Pt(0)
    tag = eyebrow.add_run("COMPANY BRIEF")
    tag.bold = True
    tag.font.size = Pt(8)
    tag.font.color.rgb = RGBColor(180, 85, 31)

    doc.add_heading(brief.company or "(unknown company)", level=0)

    meta = " · ".join(filter(None, [
        brief.domain,
        f"generated {brief.generated_at:%d %B %Y, %H:%M} UTC",
        (brief.status or "draft").upper(),
    ]))
    muted(meta, size=9)
    if brief.approved_by:
        stamp = f"approved by {brief.approved_by}"
        if brief.approved_at:
            stamp += f" on {brief.approved_at:%d %B %Y, %H:%M} UTC"
        muted(stamp, size=9)

    if brief.upcoming_meeting:
        meeting = brief.upcoming_meeting
        doc.add_heading("Upcoming meeting", level=1)
        line = doc.add_paragraph()
        run = line.add_run(f"{meeting.starts_at:%A %d %B %Y, %H:%M} — "
                           f"{meeting.title or '(untitled)'}")
        run.bold = True
        if getattr(meeting, "attendees", None):
            muted(", ".join(meeting.attendees))

    def claim_block(title: str, claims: list[BriefClaim]) -> None:
        if not claims:
            return
        doc.add_heading(title, level=1)
        for claim in claims:
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(1)
            label = paragraph.add_run(f"{_label(claim.field)}: ")
            label.bold = True
            paragraph.add_run(str(claim.value))
            note, warn = _source_note(claim)
            muted(note, size=8, italic=warn,
                  colour=(138, 90, 5) if warn else (107, 107, 102))

    claim_block("What we know", facts)
    claim_block("Recent news", news)
    claim_block("Worth raising", prep)

    if brief.talking_points:
        doc.add_heading("Talking points", level=1)
        for point in brief.talking_points:
            doc.add_paragraph(point, style="List Bullet")

    if brief.unknowns:
        doc.add_heading("Not established", level=1)
        for unknown in brief.unknowns:
            doc.add_paragraph(unknown, style="List Bullet")

    if brief.sources:
        doc.add_heading("Sources", level=1)
        for url in brief.sources:
            paragraph = doc.add_paragraph(style="List Number")
            run = paragraph.add_run(url)
            run.font.size = Pt(8)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = f"{brief.domain or brief.company} · " \
                  f"generated {brief.generated_at:%Y-%m-%d %H:%M} UTC"
    footer.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in footer.runs:
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(107, 107, 102)

    import io

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
